"""Unit tests for Template Parser."""
import pytest
import tempfile
import os
from docx import Document
from app.core.template_parser import TemplateParser
from app.exceptions import InvalidTemplateError


@pytest.fixture
def parser():
    """Provide a TemplateParser instance."""
    return TemplateParser()


@pytest.fixture
def temp_docx_with_variables():
    """Create a temporary .docx file with Jinja2 variables."""
    doc = Document()
    doc.add_paragraph("Hello {{name}}, welcome to the meeting!")
    doc.add_paragraph("Meeting date: {{date}}")
    doc.add_paragraph("Attendees:")
    doc.add_paragraph("{% for attendee in attendees %}")
    doc.add_paragraph("- {{attendee.name}} ({{attendee.email}})")
    doc.add_paragraph("{% endfor %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        yield tmp.name
    
    # Cleanup
    os.unlink(tmp.name)


@pytest.fixture
def temp_docx_without_variables():
    """Create a temporary .docx file without Jinja2 variables."""
    doc = Document()
    doc.add_paragraph("This is a plain document with no variables.")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        yield tmp.name
    
    # Cleanup
    os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_extract_variables_simple(parser, temp_docx_with_variables):
    """Test extracting variables from a document with simple and nested variables."""
    schema = await parser.extract_variables(temp_docx_with_variables)
    
    # Verify schema structure
    assert schema["type"] == "object"
    assert "properties" in schema
    assert "required" in schema
    
    # Verify simple variables
    assert "name" in schema["properties"]
    assert schema["properties"]["name"]["type"] == "string"
    assert "name" in schema["required"]
    
    assert "date" in schema["properties"]
    assert schema["properties"]["date"]["type"] == "string"
    assert "date" in schema["required"]
    
    # Verify array variable
    assert "attendees" in schema["properties"]
    assert schema["properties"]["attendees"]["type"] == "array"
    assert "attendees" in schema["required"]
    
    # Verify nested object structure (attendee.name, attendee.email)
    assert "attendee" in schema["properties"]
    assert schema["properties"]["attendee"]["type"] == "object"
    assert "properties" in schema["properties"]["attendee"]
    assert "name" in schema["properties"]["attendee"]["properties"]
    assert "email" in schema["properties"]["attendee"]["properties"]


@pytest.mark.asyncio
async def test_extract_variables_no_variables(parser, temp_docx_without_variables):
    """Test that InvalidTemplateError is raised when no variables are found."""
    with pytest.raises(InvalidTemplateError, match="No Jinja2 variables found"):
        await parser.extract_variables(temp_docx_without_variables)


@pytest.mark.asyncio
async def test_extract_variables_nested_objects(parser):
    """Test extracting nested object variables."""
    doc = Document()
    doc.add_paragraph("User: {{user.profile.name}}")
    doc.add_paragraph("Email: {{user.profile.email}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify nested structure
            assert "user" in schema["properties"]
            assert schema["properties"]["user"]["type"] == "object"
            
            user_props = schema["properties"]["user"]["properties"]
            assert "profile" in user_props
            assert user_props["profile"]["type"] == "object"
            
            profile_props = user_props["profile"]["properties"]
            assert "name" in profile_props
            assert "email" in profile_props
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_extract_variables_from_table(parser):
    """Test extracting variables from table cells."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "{{name}}"
    table.cell(1, 0).text = "Age"
    table.cell(1, 1).text = "{{age}}"
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify variables from table
            assert "name" in schema["properties"]
            assert "age" in schema["properties"]
            
        finally:
            os.unlink(tmp.name)


def test_parse_document_xml(parser, temp_docx_with_variables):
    """Test parsing document XML to extract expressions."""
    expressions = parser._parse_document_xml(temp_docx_with_variables)
    
    # Verify expressions are found
    assert len(expressions) > 0
    assert any("{{name}}" in expr for expr in expressions)
    assert any("{{date}}" in expr for expr in expressions)
    assert any("{% for attendee in attendees %}" in expr for expr in expressions)


def test_extract_expressions_from_text(parser):
    """Test extracting expressions from text."""
    text = "Hello {{name}}, your email is {{email}}. {% for item in items %}"
    expressions = parser._extract_expressions_from_text(text)
    
    assert len(expressions) == 3
    assert "{{name}}" in expressions
    assert "{{email}}" in expressions
    assert "{% for item in items %}" in expressions


def test_infer_variable_types_simple(parser):
    """Test type inference for simple variables."""
    expressions = ["{{name}}", "{{email}}"]
    schema = parser._infer_variable_types(expressions)
    
    assert schema["type"] == "object"
    assert "name" in schema["properties"]
    assert "email" in schema["properties"]
    assert schema["properties"]["name"]["type"] == "string"
    assert schema["properties"]["email"]["type"] == "string"


def test_infer_variable_types_array(parser):
    """Test type inference for array variables."""
    expressions = ["{% for item in items %}"]
    schema = parser._infer_variable_types(expressions)
    
    assert "items" in schema["properties"]
    assert schema["properties"]["items"]["type"] == "array"
    assert schema["properties"]["items"]["items"]["type"] == "object"


def test_infer_variable_types_nested(parser):
    """Test type inference for nested variables."""
    expressions = ["{{user.name}}", "{{user.email}}"]
    schema = parser._infer_variable_types(expressions)
    
    assert "user" in schema["properties"]
    assert schema["properties"]["user"]["type"] == "object"
    assert "name" in schema["properties"]["user"]["properties"]
    assert "email" in schema["properties"]["user"]["properties"]


def test_infer_variable_types_no_duplicates(parser):
    """Test that duplicate variables are not added multiple times."""
    expressions = ["{{name}}", "{{name}}", "{{name}}"]
    schema = parser._infer_variable_types(expressions)
    
    # Should only have one 'name' property
    assert "name" in schema["properties"]
    assert schema["required"].count("name") == 1
