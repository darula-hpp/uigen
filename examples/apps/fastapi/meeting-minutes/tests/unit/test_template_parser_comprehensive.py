"""Comprehensive tests for Template Parser type inference rules."""
import pytest
import tempfile
import os
from docx import Document
from app.core.template_parser import TemplateParser


@pytest.fixture
def parser():
    """Provide a TemplateParser instance."""
    return TemplateParser()


@pytest.mark.asyncio
async def test_type_inference_rule_simple_variable(parser):
    """
    Test type inference rule: {{variable}} → {"variable": {"type": "string"}}
    """
    doc = Document()
    doc.add_paragraph("{{variable}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            assert "variable" in schema["properties"]
            assert schema["properties"]["variable"]["type"] == "string"
            assert "variable" in schema["required"]
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_type_inference_rule_for_loop(parser):
    """
    Test type inference rule: 
    {% for item in items %} → {"items": {"type": "array", "items": {"type": "object"}}}
    """
    doc = Document()
    doc.add_paragraph("{% for item in items %}")
    doc.add_paragraph("{{item}}")
    doc.add_paragraph("{% endfor %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            assert "items" in schema["properties"]
            assert schema["properties"]["items"]["type"] == "array"
            assert "items" in schema["properties"]["items"]
            assert schema["properties"]["items"]["items"]["type"] == "object"
            assert "items" in schema["required"]
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_type_inference_rule_nested_object(parser):
    """
    Test type inference rule:
    {{object.property}} → {"object": {"type": "object", "properties": {"property": {"type": "string"}}}}
    """
    doc = Document()
    doc.add_paragraph("{{object.property}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            assert "object" in schema["properties"]
            assert schema["properties"]["object"]["type"] == "object"
            assert "properties" in schema["properties"]["object"]
            assert "property" in schema["properties"]["object"]["properties"]
            assert schema["properties"]["object"]["properties"]["property"]["type"] == "string"
            assert "object" in schema["required"]
            assert "property" in schema["properties"]["object"]["required"]
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_complex_nested_structure(parser):
    """Test handling of deeply nested structures."""
    doc = Document()
    doc.add_paragraph("{{company.department.manager.name}}")
    doc.add_paragraph("{{company.department.manager.email}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify company object
            assert "company" in schema["properties"]
            assert schema["properties"]["company"]["type"] == "object"
            
            # Verify department object
            company_props = schema["properties"]["company"]["properties"]
            assert "department" in company_props
            assert company_props["department"]["type"] == "object"
            
            # Verify manager object
            dept_props = company_props["department"]["properties"]
            assert "manager" in dept_props
            assert dept_props["manager"]["type"] == "object"
            
            # Verify name and email properties
            manager_props = dept_props["manager"]["properties"]
            assert "name" in manager_props
            assert "email" in manager_props
            assert manager_props["name"]["type"] == "string"
            assert manager_props["email"]["type"] == "string"
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_mixed_simple_and_nested_variables(parser):
    """Test handling of both simple and nested variables in the same document."""
    doc = Document()
    doc.add_paragraph("Title: {{title}}")
    doc.add_paragraph("Author: {{author.name}}")
    doc.add_paragraph("Email: {{author.email}}")
    doc.add_paragraph("Date: {{date}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify simple variables
            assert "title" in schema["properties"]
            assert schema["properties"]["title"]["type"] == "string"
            assert "date" in schema["properties"]
            assert schema["properties"]["date"]["type"] == "string"
            
            # Verify nested object
            assert "author" in schema["properties"]
            assert schema["properties"]["author"]["type"] == "object"
            author_props = schema["properties"]["author"]["properties"]
            assert "name" in author_props
            assert "email" in author_props
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_nested_array_in_object(parser):
    """Test handling of arrays nested within objects."""
    doc = Document()
    doc.add_paragraph("{% for item in user.items %}")
    doc.add_paragraph("{{item}}")
    doc.add_paragraph("{% endfor %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify user object
            assert "user" in schema["properties"]
            assert schema["properties"]["user"]["type"] == "object"
            
            # Verify items array within user
            user_props = schema["properties"]["user"]["properties"]
            assert "items" in user_props
            assert user_props["items"]["type"] == "array"
            assert user_props["items"]["items"]["type"] == "object"
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_multiple_arrays(parser):
    """Test handling of multiple array variables."""
    doc = Document()
    doc.add_paragraph("{% for attendee in attendees %}")
    doc.add_paragraph("{{attendee.name}}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% for task in tasks %}")
    doc.add_paragraph("{{task.title}}")
    doc.add_paragraph("{% endfor %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify both arrays
            assert "attendees" in schema["properties"]
            assert schema["properties"]["attendees"]["type"] == "array"
            assert "tasks" in schema["properties"]
            assert schema["properties"]["tasks"]["type"] == "array"
            
            # Note: 'attendee' and 'task' are loop variables and should NOT be in the schema
            # Loop variables are temporary iteration variables, not data variables
            assert "attendee" not in schema["properties"]
            assert "task" not in schema["properties"]
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_whitespace_handling(parser):
    """Test that variables with whitespace are correctly parsed."""
    doc = Document()
    doc.add_paragraph("{{  name  }}")
    doc.add_paragraph("{{   email   }}")
    doc.add_paragraph("{%  for  item  in  items  %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify variables are found despite whitespace
            assert "name" in schema["properties"]
            assert "email" in schema["properties"]
            assert "items" in schema["properties"]
            
        finally:
            os.unlink(tmp.name)


@pytest.mark.asyncio
async def test_real_world_meeting_minutes_template(parser):
    """Test with a realistic meeting minutes template."""
    doc = Document()
    doc.add_heading("Meeting Minutes", 0)
    doc.add_paragraph("Meeting Title: {{meeting.title}}")
    doc.add_paragraph("Date: {{meeting.date}}")
    doc.add_paragraph("Location: {{meeting.location}}")
    doc.add_paragraph("")
    doc.add_heading("Attendees", 1)
    doc.add_paragraph("{% for attendee in attendees %}")
    doc.add_paragraph("- {{attendee.name}} ({{attendee.role}})")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("")
    doc.add_heading("Agenda Items", 1)
    doc.add_paragraph("{% for item in agenda_items %}")
    doc.add_paragraph("{{item.number}}. {{item.title}}")
    doc.add_paragraph("   Description: {{item.description}}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("")
    doc.add_heading("Action Items", 1)
    doc.add_paragraph("{% for action in actions %}")
    doc.add_paragraph("- {{action.task}} (Assigned to: {{action.assignee}})")
    doc.add_paragraph("{% endfor %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        
        try:
            schema = await parser.extract_variables(tmp.name)
            
            # Verify meeting object
            assert "meeting" in schema["properties"]
            assert schema["properties"]["meeting"]["type"] == "object"
            meeting_props = schema["properties"]["meeting"]["properties"]
            assert "title" in meeting_props
            assert "date" in meeting_props
            assert "location" in meeting_props
            
            # Verify arrays
            assert "attendees" in schema["properties"]
            assert schema["properties"]["attendees"]["type"] == "array"
            assert "agenda_items" in schema["properties"]
            assert schema["properties"]["agenda_items"]["type"] == "array"
            assert "actions" in schema["properties"]
            assert schema["properties"]["actions"]["type"] == "array"
            
            # Note: 'attendee', 'item', and 'action' are loop variables and should NOT be in the schema
            # Loop variables are temporary iteration variables, not data variables
            assert "attendee" not in schema["properties"]
            assert "item" not in schema["properties"]
            assert "action" not in schema["properties"]
            
        finally:
            os.unlink(tmp.name)
