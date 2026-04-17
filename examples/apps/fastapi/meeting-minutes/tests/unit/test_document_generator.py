"""Unit tests for Document Generator."""
import pytest
import tempfile
import os
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
from app.core.document_generator import DocumentGenerator
from app.exceptions import RenderError


@pytest.fixture
def generator():
    """Provide a DocumentGenerator instance."""
    return DocumentGenerator()


@pytest.fixture
def temp_template_simple():
    """Create a temporary template with simple Jinja2 variables."""
    doc = Document()
    doc.add_paragraph("Hello {{name}}, welcome to the meeting!")
    doc.add_paragraph("Meeting date: {{date}}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        yield tmp.name
    
    # Cleanup
    os.unlink(tmp.name)


@pytest.fixture
def temp_template_complex():
    """Create a temporary template with complex Jinja2 structures."""
    doc = Document()
    doc.add_paragraph("Meeting: {{title}}")
    doc.add_paragraph("Date: {{date}}")
    doc.add_paragraph("Attendees:")
    doc.add_paragraph("{% for attendee in attendees %}")
    doc.add_paragraph("- {{attendee.name}} ({{attendee.email}})")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("Action Items:")
    doc.add_paragraph("{% for item in action_items %}")
    doc.add_paragraph("- {{item}}")
    doc.add_paragraph("{% endfor %}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        yield tmp.name
    
    # Cleanup
    os.unlink(tmp.name)


@pytest.fixture
def temp_output_path(tmp_path):
    """Provide a temporary output path."""
    output_path = tmp_path / "output" / "rendered.docx"
    yield str(output_path)
    
    # Cleanup
    if output_path.exists():
        output_path.unlink()


@pytest.mark.asyncio
async def test_render_template_simple(generator, temp_template_simple, temp_output_path):
    """Test rendering a template with simple variables."""
    filled_data = {
        "name": "John Doe",
        "date": "2024-01-15"
    }
    
    result_path = await generator.render_template(
        temp_template_simple,
        filled_data,
        temp_output_path
    )
    
    # Verify output file was created
    assert os.path.exists(result_path)
    assert result_path == temp_output_path
    
    # Verify content was rendered correctly
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "John Doe" in text
    assert "2024-01-15" in text
    assert "{{name}}" not in text  # Template variables should be replaced
    assert "{{date}}" not in text


@pytest.mark.asyncio
async def test_render_template_complex(generator, temp_template_complex, temp_output_path):
    """Test rendering a template with complex structures (arrays, objects)."""
    filled_data = {
        "title": "Q1 Planning Meeting",
        "date": "2024-01-15",
        "attendees": [
            {"name": "Alice Smith", "email": "alice@example.com"},
            {"name": "Bob Johnson", "email": "bob@example.com"},
            {"name": "Carol White", "email": "carol@example.com"}
        ],
        "action_items": [
            "Review budget proposal",
            "Schedule follow-up meeting",
            "Send meeting notes to team"
        ]
    }
    
    result_path = await generator.render_template(
        temp_template_complex,
        filled_data,
        temp_output_path
    )
    
    # Verify output file was created
    assert os.path.exists(result_path)
    
    # Verify content was rendered correctly
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Check title and date
    assert "Q1 Planning Meeting" in text
    assert "2024-01-15" in text
    
    # Check attendees were rendered
    assert "Alice Smith" in text
    assert "alice@example.com" in text
    assert "Bob Johnson" in text
    assert "bob@example.com" in text
    assert "Carol White" in text
    assert "carol@example.com" in text
    
    # Check action items were rendered
    assert "Review budget proposal" in text
    assert "Schedule follow-up meeting" in text
    assert "Send meeting notes to team" in text
    
    # Verify template syntax is gone
    assert "{% for" not in text
    assert "{% endfor" not in text


@pytest.mark.asyncio
async def test_render_template_empty_arrays(generator, temp_template_complex, temp_output_path):
    """Test rendering a template with empty arrays."""
    filled_data = {
        "title": "Empty Meeting",
        "date": "2024-01-15",
        "attendees": [],
        "action_items": []
    }
    
    result_path = await generator.render_template(
        temp_template_complex,
        filled_data,
        temp_output_path
    )
    
    # Verify output file was created
    assert os.path.exists(result_path)
    
    # Verify content was rendered (with empty lists)
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "Empty Meeting" in text
    assert "2024-01-15" in text


@pytest.mark.asyncio
async def test_render_template_creates_output_directory(generator, temp_template_simple, tmp_path):
    """Test that render_template creates output directory if it doesn't exist."""
    # Use a nested path that doesn't exist
    output_path = tmp_path / "nested" / "dir" / "output.docx"
    
    filled_data = {
        "name": "Test User",
        "date": "2024-01-15"
    }
    
    result_path = await generator.render_template(
        temp_template_simple,
        filled_data,
        str(output_path)
    )
    
    # Verify directory was created
    assert output_path.parent.exists()
    assert os.path.exists(result_path)


@pytest.mark.asyncio
async def test_render_template_invalid_template_path(generator, temp_output_path):
    """Test that RenderError is raised for invalid template path."""
    filled_data = {"name": "Test"}
    
    with pytest.raises(RenderError, match="Failed to render template"):
        await generator.render_template(
            "/nonexistent/template.docx",
            filled_data,
            temp_output_path
        )


@pytest.mark.asyncio
async def test_render_template_missing_variable(generator, temp_template_simple, temp_output_path):
    """Test rendering with missing variables (should use empty string or raise error)."""
    # Provide incomplete data
    filled_data = {
        "name": "John Doe"
        # Missing 'date'
    }
    
    # docxtpl may handle missing variables differently, but should not crash
    # It typically renders them as empty or raises an error
    try:
        result_path = await generator.render_template(
            temp_template_simple,
            filled_data,
            temp_output_path
        )
        # If it succeeds, verify file was created
        assert os.path.exists(result_path)
    except RenderError:
        # If it raises RenderError, that's also acceptable behavior
        pass


@pytest.mark.asyncio
async def test_render_template_with_nested_objects(generator, tmp_path):
    """Test rendering with nested object structures."""
    # Create template with nested variables
    doc = Document()
    doc.add_paragraph("User: {{user.profile.name}}")
    doc.add_paragraph("Email: {{user.profile.email}}")
    doc.add_paragraph("Department: {{user.department}}")
    
    template_path = tmp_path / "template.docx"
    doc.save(str(template_path))
    
    output_path = tmp_path / "output.docx"
    
    filled_data = {
        "user": {
            "profile": {
                "name": "Jane Doe",
                "email": "jane@example.com"
            },
            "department": "Engineering"
        }
    }
    
    result_path = await generator.render_template(
        str(template_path),
        filled_data,
        str(output_path)
    )
    
    # Verify content
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "Jane Doe" in text
    assert "jane@example.com" in text
    assert "Engineering" in text


@pytest.mark.asyncio
async def test_render_template_with_special_characters(generator, tmp_path):
    """Test rendering with special characters in data."""
    doc = Document()
    doc.add_paragraph("Title: {{title}}")
    doc.add_paragraph("Description: {{description}}")
    
    template_path = tmp_path / "template.docx"
    doc.save(str(template_path))
    
    output_path = tmp_path / "output.docx"
    
    filled_data = {
        "title": "Meeting & Planning Session",
        "description": "Review Q1 results (2024) - 100% complete!"
    }
    
    result_path = await generator.render_template(
        str(template_path),
        filled_data,
        str(output_path)
    )
    
    # Verify content is rendered (special characters may be escaped by docxtpl)
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    # Check that the main content is present (& may be escaped or removed)
    assert "Meeting" in text
    assert "Planning Session" in text
    assert "Review Q1 results (2024)" in text
    assert "100% complete!" in text


@pytest.mark.asyncio
async def test_render_template_overwrites_existing_file(generator, temp_template_simple, temp_output_path):
    """Test that rendering overwrites an existing output file."""
    # Create initial file
    filled_data_1 = {"name": "First User", "date": "2024-01-01"}
    await generator.render_template(
        temp_template_simple,
        filled_data_1,
        temp_output_path
    )
    
    # Overwrite with new data
    filled_data_2 = {"name": "Second User", "date": "2024-01-15"}
    result_path = await generator.render_template(
        temp_template_simple,
        filled_data_2,
        temp_output_path
    )
    
    # Verify new content
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "Second User" in text
    assert "2024-01-15" in text
    assert "First User" not in text


def test_render_sync_creates_directory(generator, temp_template_simple, tmp_path):
    """Test that _render_sync creates output directory."""
    output_path = tmp_path / "new_dir" / "output.docx"
    filled_data = {"name": "Test", "date": "2024-01-15"}
    
    result_path = generator._render_sync(
        temp_template_simple,
        filled_data,
        str(output_path)
    )
    
    assert os.path.exists(result_path)
    assert output_path.parent.exists()


def test_render_sync_basic_functionality(generator, temp_template_simple, tmp_path):
    """Test _render_sync basic functionality."""
    output_path = tmp_path / "output.docx"
    filled_data = {"name": "Sync Test", "date": "2024-01-15"}
    
    result_path = generator._render_sync(
        temp_template_simple,
        filled_data,
        str(output_path)
    )
    
    assert result_path == str(output_path)
    assert os.path.exists(result_path)
    
    # Verify content
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "Sync Test" in text


@pytest.mark.asyncio
async def test_render_template_with_boolean_values(generator, tmp_path):
    """Test rendering with boolean values."""
    doc = Document()
    doc.add_paragraph("Active: {{is_active}}")
    doc.add_paragraph("Completed: {{is_completed}}")
    
    template_path = tmp_path / "template.docx"
    doc.save(str(template_path))
    
    output_path = tmp_path / "output.docx"
    
    filled_data = {
        "is_active": True,
        "is_completed": False
    }
    
    result_path = await generator.render_template(
        str(template_path),
        filled_data,
        str(output_path)
    )
    
    # Verify boolean values are rendered
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "True" in text or "true" in text.lower()
    assert "False" in text or "false" in text.lower()


@pytest.mark.asyncio
async def test_render_template_with_numeric_values(generator, tmp_path):
    """Test rendering with numeric values."""
    doc = Document()
    doc.add_paragraph("Count: {{count}}")
    doc.add_paragraph("Price: {{price}}")
    
    template_path = tmp_path / "template.docx"
    doc.save(str(template_path))
    
    output_path = tmp_path / "output.docx"
    
    filled_data = {
        "count": 42,
        "price": 99.99
    }
    
    result_path = await generator.render_template(
        str(template_path),
        filled_data,
        str(output_path)
    )
    
    # Verify numeric values are rendered
    doc = Document(result_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    assert "42" in text
    assert "99.99" in text
