"""Template parser for extracting Jinja2 variables from Word documents."""
import re
from typing import Dict, Any, List, Set
from docx import Document
from app.exceptions import InvalidTemplateError
from app.schemas import JinjaShape


class TemplateParser:
    """Parser for extracting Jinja2 variables from .docx templates."""
    
    # Regex patterns for Jinja2 syntax
    VARIABLE_PATTERN = r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_\.]*)\s*\}\}'
    FOR_LOOP_PATTERN = r'\{%\s*for\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+([a-zA-Z_][a-zA-Z0-9_\.]*)\s*%\}'
    
    async def extract_variables(self, docx_path: str) -> JinjaShape:
        """
        Extract all Jinja2 variables from a Word document.
        
        Args:
            docx_path: Path to the .docx file
            
        Returns:
            JinjaShape: JSON schema describing variable names and types
            
        Raises:
            InvalidTemplateError: If no variables found or file is invalid
        """
        try:
            # Parse document to extract raw Jinja2 expressions
            variables = self._parse_document_xml(docx_path)
            
            if not variables:
                raise InvalidTemplateError("No Jinja2 variables found in template")
            
            # Infer types and generate JSON schema
            jinja_shape = self._infer_variable_types(variables)
            
            return jinja_shape
            
        except Exception as e:
            if isinstance(e, InvalidTemplateError):
                raise
            raise InvalidTemplateError(f"Failed to parse template: {str(e)}")
    
    def _parse_document_xml(self, docx_path: str) -> List[str]:
        """
        Extract raw Jinja2 expressions from document XML.
        
        Args:
            docx_path: Path to the .docx file
            
        Returns:
            List of Jinja2 expressions found in the document
        """
        doc = Document(docx_path)
        expressions = []
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            expressions.extend(self._extract_expressions_from_text(text))
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    expressions.extend(self._extract_expressions_from_text(text))
        
        return expressions
    
    def _extract_expressions_from_text(self, text: str) -> List[str]:
        """
        Extract Jinja2 expressions from a text string.
        
        Args:
            text: Text to search for Jinja2 expressions
            
        Returns:
            List of expressions found
        """
        expressions = []
        
        # Find variable expressions: {{variable}}
        variable_matches = re.finditer(self.VARIABLE_PATTERN, text)
        for match in variable_matches:
            expressions.append(f"{{{{{match.group(1)}}}}}")
        
        # Find for loop expressions: {% for item in items %}
        for_loop_matches = re.finditer(self.FOR_LOOP_PATTERN, text)
        for match in for_loop_matches:
            expressions.append(f"{{% for {match.group(1)} in {match.group(2)} %}}")
        
        return expressions
    
    def _infer_variable_types(self, expressions: List[str]) -> Dict[str, Any]:
        """
        Infer types from Jinja2 variable syntax and generate JSON schema.
        
        Args:
            expressions: List of Jinja2 expressions
            
        Returns:
            JSON schema describing the variables
        """
        schema = {
            "type": "object",
            "properties": {},
            "required": []
        }
        
        # Track which variables we've seen to avoid duplicates
        seen_variables: Set[str] = set()
        
        # Track loop variables to exclude them from the schema
        loop_variables: Set[str] = set()
        
        # First pass: identify all loop variables
        for expression in expressions:
            for_loop_match = re.match(self.FOR_LOOP_PATTERN, expression)
            if for_loop_match:
                loop_var = for_loop_match.group(1)
                loop_variables.add(loop_var)
        
        # Second pass: process all expressions
        for expression in expressions:
            # Handle variable expressions: {{variable}}
            variable_match = re.match(self.VARIABLE_PATTERN, expression)
            if variable_match:
                var_path = variable_match.group(1)
                # Skip if this is a loop variable or starts with a loop variable
                root_var = var_path.split('.')[0]
                if root_var not in loop_variables:
                    self._add_variable_to_schema(schema, var_path, seen_variables)
            
            # Handle for loop expressions: {% for item in items %}
            for_loop_match = re.match(self.FOR_LOOP_PATTERN, expression)
            if for_loop_match:
                loop_var = for_loop_match.group(1)
                collection_var = for_loop_match.group(2)
                self._add_array_to_schema(schema, collection_var, seen_variables)
        
        return schema
    
    def _add_variable_to_schema(
        self, 
        schema: Dict[str, Any], 
        var_path: str, 
        seen_variables: Set[str]
    ) -> None:
        """
        Add a variable to the JSON schema, handling nested paths.
        
        Args:
            schema: The schema to modify
            var_path: Variable path (e.g., "user.name")
            seen_variables: Set of variables already processed
        """
        parts = var_path.split('.')
        
        if len(parts) == 1:
            # Simple variable: {{variable}}
            var_name = parts[0]
            if var_name not in seen_variables:
                schema["properties"][var_name] = {"type": "string"}
                schema["required"].append(var_name)
                seen_variables.add(var_name)
        else:
            # Nested variable: {{object.property}}
            root_var = parts[0]
            
            # Create object structure if it doesn't exist, or upgrade from simple type
            if root_var not in schema["properties"]:
                schema["properties"][root_var] = {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
                if root_var not in schema["required"]:
                    schema["required"].append(root_var)
            elif schema["properties"][root_var].get("type") != "object":
                # Upgrade from simple type to object type
                schema["properties"][root_var] = {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
            
            # Navigate/create nested structure
            current = schema["properties"][root_var]
            for i, part in enumerate(parts[1:], 1):
                is_last = i == len(parts) - 1
                
                if is_last:
                    # Leaf property
                    if part not in current["properties"]:
                        current["properties"][part] = {"type": "string"}
                        current["required"].append(part)
                        seen_variables.add(var_path)
                else:
                    # Intermediate object
                    if part not in current["properties"]:
                        current["properties"][part] = {
                            "type": "object",
                            "properties": {},
                            "required": []
                        }
                        current["required"].append(part)
                    elif current["properties"][part].get("type") != "object":
                        # Upgrade from simple/array type to object type
                        current["properties"][part] = {
                            "type": "object",
                            "properties": {},
                            "required": []
                        }
                    current = current["properties"][part]
    
    def _add_array_to_schema(
        self, 
        schema: Dict[str, Any], 
        collection_path: str, 
        seen_variables: Set[str]
    ) -> None:
        """
        Add an array variable to the JSON schema.
        
        Args:
            schema: The schema to modify
            collection_path: Collection variable path (e.g., "items")
            seen_variables: Set of variables already processed
        """
        parts = collection_path.split('.')
        
        if len(parts) == 1:
            # Simple array: {% for item in items %}
            var_name = parts[0]
            
            # Check if variable already exists with a different type and needs upgrading
            if var_name in schema["properties"] and schema["properties"][var_name].get("type") != "array":
                # Upgrade to array type
                schema["properties"][var_name] = {
                    "type": "array",
                    "items": {"type": "object"}
                }
                seen_variables.add(var_name)
            elif var_name not in seen_variables:
                # Add new array variable
                schema["properties"][var_name] = {
                    "type": "array",
                    "items": {"type": "object"}
                }
                schema["required"].append(var_name)
                seen_variables.add(var_name)
        else:
            # Nested array: {% for item in user.items %}
            root_var = parts[0]
            
            # Create object structure if it doesn't exist, or upgrade from simple type
            if root_var not in schema["properties"]:
                schema["properties"][root_var] = {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
                if root_var not in schema["required"]:
                    schema["required"].append(root_var)
            elif schema["properties"][root_var].get("type") != "object":
                # Upgrade from simple type to object type
                schema["properties"][root_var] = {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
            
            # Navigate/create nested structure
            current = schema["properties"][root_var]
            for i, part in enumerate(parts[1:], 1):
                is_last = i == len(parts) - 1
                
                if is_last:
                    # Leaf array property
                    if part not in current["properties"]:
                        current["properties"][part] = {
                            "type": "array",
                            "items": {"type": "object"}
                        }
                        current["required"].append(part)
                        seen_variables.add(collection_path)
                else:
                    # Intermediate object
                    if part not in current["properties"]:
                        current["properties"][part] = {
                            "type": "object",
                            "properties": {},
                            "required": []
                        }
                        current["required"].append(part)
                    elif current["properties"][part].get("type") != "object":
                        # Upgrade from simple/array type to object type
                        current["properties"][part] = {
                            "type": "object",
                            "properties": {},
                            "required": []
                        }
                    current = current["properties"][part]
